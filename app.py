from flask import Flask, render_template, request, jsonify, session
import os, sys, tempfile, base64, re
import nltk, PyPDF2, docx, fitz
from io import BytesIO
from PIL import Image
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter as pdf2docx_Converter
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize
from nltk.sentiment import SentimentIntensityAnalyzer

# Optional: COM for docx2pdf on Windows
if sys.platform.startswith("win"):
    try:
        import pythoncom
    except Exception:
        pythoncom = None
else:
    pythoncom = None

# ---------- NLTK bootstrap (robust) ----------
def _safe_nltk_download(pkg: str):
    try:
        nltk.download(pkg, quiet=True)
        return True
    except Exception:
        return False

# Newer NLTK sometimes requires "punkt_tab" in addition to "punkt"
for pkg in ("punkt", "punkt_tab", "stopwords", "vader_lexicon"):
    _safe_nltk_download(pkg)

def _has_nltk_tokenizer():
    try:
        nltk.data.find("tokenizers/punkt")
        return True
    except LookupError:
        return False

def _has_punkt_tab():
    try:
        nltk.data.find("tokenizers/punkt_tab")
        return True
    except LookupError:
        return False

def safe_sent_tokenize(text: str):
    """
    Try NLTK sentence tokenization; fallback to a regex-based splitter if unavailable.
    """
    text = (text or "").strip()
    if not text:
        return []
    try:
        if not _has_nltk_tokenizer():
            _safe_nltk_download("punkt")
        if not _has_punkt_tab():
            _safe_nltk_download("punkt_tab")
        return sent_tokenize(text)
    except Exception:
        # Regex fallback: split on ., !, ? followed by whitespace
        chunks = re.split(r'(?<=[.!?])\s+(?=[A-Z0-9"\'])', text)
        if len(chunks) == 1:
            chunks = re.split(r'(?<=[.!?])\s+', text)
        return [c.strip() for c in chunks if c.strip()]

def safe_stopwords():
    """
    Return NLTK English stopwords, or a small fallback set if unavailable.
    """
    try:
        return set(stopwords.words("english"))
    except Exception:
        _safe_nltk_download("stopwords")
        try:
            return set(stopwords.words("english"))
        except Exception:
            return {
                "the","a","an","is","are","was","were","be","been","being",
                "and","or","but","if","then","than","that","this","to","for",
                "of","in","on","at","by","with","as","from","it","its","into",
                "about","over","under","again","further","more","most","so",
            }

# ---------- Flask app ----------
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "change_me_in_prod")

# Initialize sentiment analyzer if possible
try:
    sid = SentimentIntensityAnalyzer()
except Exception as e:
    print(f"[WARN] Sentiment init failed: {e}")
    sid = None

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file or file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ['.txt', '.pdf', '.docx']:
        return jsonify({'error': 'Unsupported file type'}), 400

    # Save temp file and close handle so other libs can read it
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_ext)
    try:
        file.save(temp_file.name)
    finally:
        try:
            temp_file.close()
        except Exception:
            pass

    session['uploaded_file'] = temp_file.name
    session['file_ext'] = file_ext

    original_text = ""
    preview_type = 'text'
    preview_content = ""

    try:
        if file_ext == '.txt':
            with open(temp_file.name, 'r', encoding='utf-8', errors='ignore') as f:
                original_text = f.read()
            preview_content = original_text

        elif file_ext == '.pdf':
            # Text extraction (best-effort; image-only PDFs may yield empty text)
            try:
                with open(temp_file.name, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        t = page.extract_text() or ""
                        if t:
                            original_text += t + "\n"
            except Exception:
                original_text = (original_text or "").strip()

            # Render first few pages as PNGs for preview
            images = []
            try:
                doc_pdf = fitz.open(temp_file.name)
                for page_num in range(min(len(doc_pdf), 5)):
                    page = doc_pdf.load_page(page_num)
                    pix = page.get_pixmap(alpha=False, dpi=144)
                    png_bytes = pix.tobytes("png")
                    # Downscale for transport
                    img = Image.open(BytesIO(png_bytes))
                    img.thumbnail((980, 1200), Image.Resampling.LANCZOS)
                    buffered = BytesIO()
                    img.save(buffered, format="PNG")
                    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                    images.append(img_base64)
                doc_pdf.close()
            except Exception:
                # If rendering fails, leave images empty; frontend can handle
                images = []
            preview_type = 'image'
            preview_content = images

        elif file_ext == '.docx':
            # Extract text (simple preview as text)
            docx_doc = docx.Document(temp_file.name)
            original_text = "\n".join(p.text for p in docx_doc.paragraphs)
            preview_content = original_text

        return jsonify({
            'preview_type': preview_type,
            'preview_content': preview_content,
            'text': (original_text or "").strip(),
            'filename': file.filename
        })
    except Exception as e:
        return jsonify({'error': f'Failed to process file: {str(e)}'}), 500

@app.route('/convert', methods=['POST'])
def convert_file():
    # Use the file already uploaded (stored in session by /upload)
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded yet'}), 400
    temp_file_name = session['uploaded_file']
    file_ext = session.get('file_ext', '').lower()

    output_format = request.form.get('format')
    if not output_format:
        return jsonify({'error': 'No output format specified'}), 400
    output_format = output_format.lower().lstrip(".")

    try:
        # Create a temp output file
        output_file = tempfile.NamedTemporaryFile(delete=False, suffix=f'.{output_format}')
        output_path = output_file.name
        output_file.close()

        if file_ext == '.docx' and output_format == 'pdf':
            # docx2pdf typically needs MS Word on Windows
            if sys.platform.startswith("win") and pythoncom is not None:
                try:
                    pythoncom.CoInitialize()
                except Exception:
                    pass
            try:
                docx2pdf_convert(temp_file_name, output_path)
            finally:
                if sys.platform.startswith("win") and pythoncom is not None:
                    try:
                        pythoncom.CoUninitialize()
                    except Exception:
                        pass

        elif file_ext == '.pdf' and output_format == 'docx':
            cv = pdf2docx_Converter(temp_file_name)
            cv.convert(output_path)
            cv.close()

        elif output_format == 'txt':
            # Normalize any input to TXT
            if file_ext == '.txt':
                with open(temp_file_name, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            elif file_ext == '.pdf':
                with open(temp_file_name, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        t = page.extract_text() or ""
                        if t:
                            text += t + "\n"
            elif file_ext == '.docx':
                docx_doc = docx.Document(temp_file_name)
                text = "\n".join(p.text for p in docx_doc.paragraphs)
            else:
                return jsonify({'error': 'Unsupported input for TXT conversion'}), 400

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)

        elif file_ext == '.txt' and output_format == 'docx':
            # Simple txt -> docx conversion
            with open(temp_file_name, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            from docx import Document as DocxDoc
            docx_out = DocxDoc()
            for line in text.splitlines():
                docx_out.add_paragraph(line)
            docx_out.save(output_path)

        elif file_ext == '.txt' and output_format == 'pdf':
            # Basic txt -> pdf by writing text into a 1-page PDF using PyMuPDF
            text = open(temp_file_name, 'r', encoding='utf-8', errors='ignore').read()
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            rect = fitz.Rect(40, 40, 560, 800)
            page.insert_textbox(rect, text)
            pdf_doc.save(output_path)
            pdf_doc.close()

        else:
            return jsonify({'error': 'Unsupported conversion'}), 400

        # Return the converted file as base64 in JSON
        with open(output_path, 'rb') as f:
            file_data = base64.b64encode(f.read()).decode('utf-8')
        try:
            os.remove(output_path)
        except Exception:
            pass

        return jsonify({
            'filename': f'converted_file.{output_format}',
            'file_data': file_data
        })
    except Exception as e:
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

@app.route('/summarize', methods=['POST'])
def summarize():
    # Frontend sends { text: "...", num_sentences?: N }
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded yet'}), 400

    data = request.get_json(silent=True) or {}
    text = data.get('text', '')
    num_sentences = data.get('num_sentences', 3)

    if not isinstance(num_sentences, int):
        try:
            num_sentences = int(num_sentences)
        except Exception:
            num_sentences = 3
    num_sentences = max(1, min(10, num_sentences))  # keep it reasonable

    if not (text or "").strip():
        return jsonify({'error': 'No text provided'}), 400

    sentences = safe_sent_tokenize(text)
    if not sentences:
        return jsonify({'summary': ''})
    if len(sentences) <= num_sentences:
        return jsonify({'summary': text})

    sw = safe_stopwords()

    # Score sentences: content words count, lightly boosted for earlier sentences
    scored = []
    total = len(sentences)
    for i, sent in enumerate(sentences):
        words = [w.lower().strip(",.!?;:()[]\"'") for w in sent.split()]
        content = [w for w in words if w and w not in sw]
        # earlier sentences get a small positional boost (front-loading)
        position_boost = 1.0 + 0.25 * (1 - i / max(1, total - 1))
        score = len(content) * position_boost
        scored.append((i, sent, score))

    scored.sort(key=lambda x: x[2], reverse=True)
    k = min(num_sentences, len(scored))
    selected_idxs = sorted([idx for idx, _, _ in scored[:k]])
    summary = " ".join([sentences[i] for i in selected_idxs])

    return jsonify({'summary': summary})

@app.route('/emotion', methods=['POST'])
def emotion_analysis():
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded yet'}), 400

    data = request.get_json(silent=True) or {}
    text = data.get('text', '')
    if not (text or "").strip():
        return jsonify({'error': 'No text provided'}), 400

    if sid is None:
        return jsonify({'error': 'Sentiment analyzer not initialized'}), 500

    scores = sid.polarity_scores(text)
    return jsonify({
        'positive': f"{scores.get('pos', 0.0)*100:.1f}%",
        'negative': f"{scores.get('neg', 0.0)*100:.1f}%",
        'neutral': f"{scores.get('neu', 0.0)*100:.1f}%",
        'compound': f"{scores.get('compound', 0.0):.3f}"
    })

if __name__ == '__main__':
    # Use host='0.0.0.0' if you want to access from other devices in your network
    app.run(debug=True)